"""
Text Extraction Module - Stage 3
Extracts text from various document formats (PDF, HWP, Excel, Word, etc.)
"""

import asyncio
import asyncpg
from pathlib import Path
from typing import Optional, Dict, List
import logging
import os

# PDF extraction
try:
    import PyPDF2
    import pdfplumber
    HAS_PDF = True
except ImportError:
    HAS_PDF = False
    logging.warning("PDF libraries not installed. Install: pip install PyPDF2 pdfplumber")

# HWP extraction
try:
    import olefile
    HAS_HWP = True
except ImportError:
    HAS_HWP = False
    logging.warning("HWP library not installed. Install: pip install olefile")

# Excel extraction
try:
    import openpyxl
    import xlrd
    HAS_EXCEL = True
except ImportError:
    HAS_EXCEL = False
    logging.warning("Excel libraries not installed. Install: pip install openpyxl xlrd")

# Word extraction
try:
    import docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    logging.warning("Word library not installed. Install: pip install python-docx")

logger = logging.getLogger(__name__)


class TextExtractor:
    """Extract text from various document formats"""

    # def __init__(self):
        # self.data_lake_root = Path(data_lake_root)

    # ========================================================================
    # HWP/HWPX Encryption Detection
    # ========================================================================

    def _is_hwp_encrypted(self, file_path: Path) -> dict:
        """
        Quick check if HWP file is encrypted without full parsing.
        Returns dict with encryption status and type.
        """
        try:
            import olefile
            if not olefile.isOleFile(str(file_path)):
                return {'encrypted': False, 'error': 'Not a valid OLE file'}

            ole = olefile.OleFileIO(str(file_path))

            # Read FileHeader
            if not ole.exists('FileHeader'):
                ole.close()
                return {'encrypted': False, 'error': 'No FileHeader stream'}

            header_data = ole.openstream('FileHeader').read()
            ole.close()

            if len(header_data) >= 40:
                # Parse properties flags at offset 36-40
                properties = int.from_bytes(header_data[36:40], 'little')

                encrypted = bool(properties & 0x02)        # Standard encryption
                distributed = bool(properties & 0x04)      # Distribution protection (배포용)
                drm = bool(properties & 0x10)              # DRM protection
                cert_encrypted = bool(properties & 0x100)  # Certificate encryption
                cert_drm = bool(properties & 0x200)        # Certificate DRM

                if encrypted or cert_encrypted:
                    return {
                        'encrypted': True,
                        'encryption_type': 'password',
                        'details': 'Password-protected HWP file'
                    }
                elif drm or cert_drm:
                    return {
                        'encrypted': True,
                        'encryption_type': 'drm',
                        'details': 'DRM-protected HWP file'
                    }
                elif distributed:
                    # Distribution docs use XOR encoding, may be extractable
                    return {
                        'encrypted': False,
                        'distributed': True,
                        'encryption_type': 'distribution',
                        'details': 'Distribution-protected document (may be extractable)'
                    }

            return {'encrypted': False}

        except Exception as e:
            return {'encrypted': False, 'error': str(e)}

    def _is_hwpx_encrypted(self, file_path: Path) -> dict:
        """
        Quick check if HWPX file is encrypted.
        Returns dict with encryption status and type.
        """
        try:
            import zipfile

            if not zipfile.is_zipfile(str(file_path)):
                return {'encrypted': False, 'error': 'Not a valid ZIP file'}

            with zipfile.ZipFile(str(file_path), 'r') as zf:
                file_list = zf.namelist()

                # Check for encryption markers
                if 'Contents/EncryptedPackage' in file_list:
                    return {
                        'encrypted': True,
                        'encryption_type': 'package',
                        'details': 'HWPX file has encrypted package'
                    }

                # Check version.xml for encryption info
                if 'version.xml' in file_list:
                    try:
                        content = zf.read('version.xml').decode('utf-8', errors='replace')
                        if 'encrypted="true"' in content.lower() or 'encrypt' in content.lower():
                            return {
                                'encrypted': True,
                                'encryption_type': 'standard',
                                'details': 'HWPX file marked as encrypted'
                            }
                    except:
                        pass

                # Check if section files are missing (could indicate encryption)
                if 'Contents/section0.xml' not in file_list:
                    # Could be encrypted or corrupted
                    if 'Contents/header.xml' in file_list:
                        return {
                            'encrypted': True,
                            'encryption_type': 'unknown',
                            'details': 'HWPX file missing section content (possibly encrypted)'
                        }

            return {'encrypted': False}

        except zipfile.BadZipFile:
            return {'encrypted': False, 'error': 'Corrupted or invalid HWPX file'}
        except Exception as e:
            return {'encrypted': False, 'error': str(e)}

    # ========================================================================
    # PDF Extraction
    # ========================================================================

    def extract_pdf_text(self, file_path: Path) -> Dict[str, any]:
        """
        Extract text from PDF using both PyPDF2 and pdfplumber
        Returns: {
            'text': str,
            'page_count': int,
            'method': str,
            'sections': List[Dict]
        }
        """
        if not HAS_PDF:
            return {
                'text': '',
                'error': 'PDF libraries not installed',
                'page_count': 0
            }

        try:
            # Try pdfplumber first (better for tables and layout)
            text_parts = []
            sections = []

            with pdfplumber.open(file_path) as pdf:
                page_count = len(pdf.pages)

                for page_num, page in enumerate(pdf.pages, 1):
                    page_text = page.extract_text()

                    if page_text:
                        text_parts.append(page_text)

                        # Create section for each page
                        sections.append({
                            'section_type': 'page',
                            'section_title': f'Page {page_num}',
                            'text': page_text,
                            'page_number': page_num,
                            'section_order': page_num,
                            'word_count': len(page_text.split()),
                            'char_count': len(page_text)
                        })

            full_text = '\n\n'.join(text_parts)

            return {
                'text': full_text,
                'page_count': page_count,
                'method': 'pdfplumber',
                'sections': sections,
                'word_count': len(full_text.split()),
                'char_count': len(full_text)
            }

        except Exception as e:
            logger.warning(f"pdfplumber failed for {file_path.name}, trying PyPDF2: {e}")

            # Fallback to PyPDF2
            try:
                with open(file_path, 'rb') as f:
                    pdf_reader = PyPDF2.PdfReader(f)
                    page_count = len(pdf_reader.pages)
                    text_parts = []
                    sections = []

                    for page_num, page in enumerate(pdf_reader.pages, 1):
                        page_text = page.extract_text()

                        if page_text:
                            text_parts.append(page_text)

                            sections.append({
                                'section_type': 'page',
                                'section_title': f'Page {page_num}',
                                'text': page_text,
                                'page_number': page_num,
                                'section_order': page_num,
                                'word_count': len(page_text.split()),
                                'char_count': len(page_text)
                            })

                    full_text = '\n\n'.join(text_parts)

                    return {
                        'text': full_text,
                        'page_count': page_count,
                        'method': 'PyPDF2',
                        'sections': sections,
                        'word_count': len(full_text.split()),
                        'char_count': len(full_text)
                    }

            except Exception as e2:
                logger.error(f"Both PDF methods failed for {file_path.name}: {e2}")
                return {
                    'text': '',
                    'error': str(e2),
                    'page_count': 0
                }

    # ========================================================================
    # HWP Extraction
    # ========================================================================

    def extract_hwp_text(self, file_path: Path) -> Dict[str, any]:
        """
        Extract text from HWP (한글) files using proper HWP parser.
        Checks for encryption before attempting extraction.
        Returns: {
            'text': str,
            'method': str,
            'sections': List[Dict],
            'encrypted': bool (if applicable),
            'encryption_type': str (if applicable)
        }
        """
        if not HAS_HWP:
            return {
                'text': '',
                'error': 'HWP library not installed. Install: pip install olefile'
            }

        # Check for encryption FIRST before attempting full parse
        enc_check = self._is_hwp_encrypted(file_path)
        if enc_check.get('encrypted'):
            logger.warning(f"  ⚠️  HWP file is encrypted: {enc_check.get('details', 'Unknown encryption')}")
            return {
                'text': '',
                'error': f"HWP file is encrypted ({enc_check.get('encryption_type', 'unknown')})",
                'encrypted': True,
                'encryption_type': enc_check.get('encryption_type', 'unknown'),
                'sections': []
            }

        # Log if distribution-protected (still attempt extraction)
        if enc_check.get('distributed'):
            logger.info(f"  ℹ️  Distribution-protected document, attempting extraction...")

        try:
            from .hwp_extraction import Document

            # Use the proper HWP parser
            doc = Document.read_hwp_document(str(file_path), verbose=False)

            if doc is None:
                # Check if it failed due to encryption we didn't detect
                return {
                    'text': '',
                    'error': 'Failed to parse HWP file (may be encrypted or corrupted)',
                    'sections': []
                }

            # Check file header for encryption flag from parsed doc
            if doc.file_header.get('encrypted') or doc.file_header.get('cert_encrypted'):
                return {
                    'text': '',
                    'error': 'HWP file is encrypted (detected during parsing)',
                    'encrypted': True,
                    'encryption_type': 'password',
                    'sections': []
                }

            # Extract clean text with Chinese and control character removal
            full_text = doc.clean_text(remove_chinese=True, remove_control=True)

            # Get sections data
            sections = doc.get_sections_data()

            if not full_text or not sections:
                return {
                    'text': '',
                    'error': 'No text found in HWP file (may be image-based or encrypted)',
                    'sections': []
                }

            logger.info(f"  [OK] Extracted {len(full_text)} chars from HWP")

            return {
                'text': full_text,
                'method': 'hwp_extraction (proper parser)',
                'sections': sections,
                'word_count': len(full_text.split()) if full_text else 0,
                'char_count': len(full_text) if full_text else 0
            }

        except Exception as e:
            error_msg = str(e).lower()
            # Check if error indicates encryption
            if 'decrypt' in error_msg or 'encrypt' in error_msg or 'password' in error_msg:
                return {
                    'text': '',
                    'error': f'HWP file appears to be encrypted: {e}',
                    'encrypted': True,
                    'encryption_type': 'unknown',
                    'sections': []
                }
            logger.error(f"HWP extraction failed for {file_path.name}: {e}")
            return {
                'text': '',
                'error': str(e)
            }

    # ========================================================================
    # HWPX Extraction
    # ========================================================================

    def extract_hwpx_text(self, file_path: Path) -> Dict[str, any]:
        """
        Extract text from HWPX (한글) files using XML parser.
        Checks for encryption before attempting extraction.
        Returns: {
            'text': str,
            'method': str,
            'sections': List[Dict],
            'encrypted': bool (if applicable),
            'encryption_type': str (if applicable)
        }
        """
        # Check for encryption FIRST before attempting full parse
        enc_check = self._is_hwpx_encrypted(file_path)
        if enc_check.get('encrypted'):
            logger.warning(f"  ⚠️  HWPX file is encrypted: {enc_check.get('details', 'Unknown encryption')}")
            return {
                'text': '',
                'error': f"HWPX file is encrypted ({enc_check.get('encryption_type', 'unknown')})",
                'encrypted': True,
                'encryption_type': enc_check.get('encryption_type', 'unknown'),
                'sections': []
            }

        try:
            from .hwpx_extraction import Document

            # Use the proper HWPX parser
            doc = Document.read_hwpx_document(str(file_path), verbose=False)

            if doc is None:
                return {
                    'text': '',
                    'error': 'Failed to parse HWPX file (may be encrypted or corrupted)',
                    'sections': []
                }

            # Extract clean text with Chinese and control character removal
            full_text = doc.clean_text(remove_chinese=True, remove_control=True)

            # Get sections data
            sections = doc.get_sections_data()

            if not full_text or not sections:
                return {
                    'text': '',
                    'error': 'No text found in HWPX file (may be image-based or encrypted)',
                    'sections': []
                }

            logger.info(f"  [OK] Extracted {len(full_text)} chars from HWPX")

            return {
                'text': full_text,
                'method': 'hwpx_extraction (XML parser)',
                'sections': sections,
                'word_count': len(full_text.split()) if full_text else 0,
                'char_count': len(full_text) if full_text else 0
            }

        except Exception as e:
            error_msg = str(e).lower()
            # Check if error indicates encryption
            if 'decrypt' in error_msg or 'encrypt' in error_msg or 'password' in error_msg:
                return {
                    'text': '',
                    'error': f'HWPX file appears to be encrypted: {e}',
                    'encrypted': True,
                    'encryption_type': 'unknown',
                    'sections': []
                }
            logger.error(f"HWPX extraction failed for {file_path.name}: {e}")
            return {
                'text': '',
                'error': str(e)
            }

    # ========================================================================
    # Excel Extraction
    # ========================================================================

    def extract_excel_text(self, file_path: Path) -> Dict[str, any]:
        """Extract text from Excel files (XLSX, XLS)"""
        if not HAS_EXCEL:
            return {
                'text': '',
                'error': 'Excel libraries not installed'
            }

        try:
            text_parts = []
            sections = []

            # Try openpyxl for .xlsx
            if file_path.suffix.lower() in ['.xlsx', '.xlsm']:
                try:
                    workbook = openpyxl.load_workbook(file_path, data_only=True)
                except Exception as load_error:
                    # Try without data_only mode as fallback
                    try:
                        workbook = openpyxl.load_workbook(file_path, data_only=False)
                    except Exception as e:
                        logger.error(f"Excel file corrupted or password-protected: {file_path.name}: {e}")
                        return {
                            'text': '',
                            'error': f'Failed to load Excel file: {e}'
                        }

                for sheet_num, sheet_name in enumerate(workbook.sheetnames, 1):
                    try:
                        sheet = workbook[sheet_name]
                        sheet_text_parts = []

                        # Use try-except for each row to handle malformed cells
                        try:
                            for row in sheet.iter_rows(values_only=True):
                                try:
                                    row_text = '\t'.join(str(cell) if cell is not None else '' for cell in row)
                                    if row_text.strip():
                                        sheet_text_parts.append(row_text)
                                except Exception as row_error:
                                    # Skip malformed rows
                                    continue
                        except Exception as iter_error:
                            # If iter_rows fails, try alternative method
                            logger.warning(f"  ⚠️  Sheet {sheet_name} has issues, trying alternate extraction")
                            try:
                                max_row = sheet.max_row or 1
                                max_col = sheet.max_column or 1
                                for row_idx in range(1, min(max_row + 1, 10000)):  # Limit to 10k rows
                                    row_cells = []
                                    for col_idx in range(1, min(max_col + 1, 100)):  # Limit to 100 cols
                                        try:
                                            cell_value = sheet.cell(row=row_idx, column=col_idx).value
                                            row_cells.append(str(cell_value) if cell_value else '')
                                        except:
                                            row_cells.append('')
                                    row_text = '\t'.join(row_cells)
                                    if row_text.strip():
                                        sheet_text_parts.append(row_text)
                            except Exception as alt_error:
                                logger.warning(f"  ⚠️  Could not extract sheet {sheet_name}: {alt_error}")
                                continue

                        sheet_text = '\n'.join(sheet_text_parts)
                        if sheet_text.strip():  # Only add non-empty sheets
                            text_parts.append(f"Sheet: {sheet_name}\n{sheet_text}")

                            sections.append({
                                'section_type': 'sheet',
                                'section_title': sheet_name,
                                'text': sheet_text,
                                'section_order': sheet_num,
                                'word_count': len(sheet_text.split()),
                                'char_count': len(sheet_text)
                            })

                    except Exception as sheet_error:
                        logger.warning(f"  ⚠️  Skipping sheet {sheet_name}: {sheet_error}")
                        continue

            # Try xlrd for .xls
            elif file_path.suffix.lower() == '.xls':
                try:
                    workbook = xlrd.open_workbook(file_path)
                except Exception as e:
                    logger.error(f"XLS file corrupted or unsupported: {file_path.name}: {e}")
                    return {
                        'text': '',
                        'error': f'Failed to load XLS file: {e}'
                    }

                for sheet_num in range(workbook.nsheets):
                    try:
                        sheet = workbook.sheet_by_index(sheet_num)
                        sheet_name = sheet.name
                        sheet_text_parts = []

                        for row_idx in range(sheet.nrows):
                            try:
                                row = sheet.row(row_idx)
                                row_text = '\t'.join(str(cell.value) for cell in row)
                                if row_text.strip():
                                    sheet_text_parts.append(row_text)
                            except Exception as row_error:
                                continue

                        sheet_text = '\n'.join(sheet_text_parts)
                        if sheet_text.strip():
                            text_parts.append(f"Sheet: {sheet_name}\n{sheet_text}")

                            sections.append({
                                'section_type': 'sheet',
                                'section_title': sheet_name,
                                'text': sheet_text,
                                'section_order': sheet_num + 1,
                                'word_count': len(sheet_text.split()),
                                'char_count': len(sheet_text)
                            })

                    except Exception as sheet_error:
                        logger.warning(f"  ⚠️  Skipping sheet index {sheet_num}: {sheet_error}")
                        continue

            full_text = '\n\n'.join(text_parts)

            if not full_text.strip():
                return {
                    'text': '',
                    'error': 'No text content found in Excel file (may be empty or image-based)',
                    'sections': []
                }

            return {
                'text': full_text,
                'method': 'openpyxl/xlrd',
                'sections': sections,
                'word_count': len(full_text.split()),
                'char_count': len(full_text)
            }

        except Exception as e:
            logger.error(f"Excel extraction failed for {file_path.name}: {e}")
            return {
                'text': '',
                'error': str(e)
            }

    # ========================================================================
    # Word Extraction
    # ========================================================================

    def extract_word_text(self, file_path: Path) -> Dict[str, any]:
        """Extract text from Word documents (DOCX)"""
        if not HAS_DOCX:
            return {
                'text': '',
                'error': 'Word library not installed'
            }

        try:
            doc = docx.Document(file_path)
            text_parts = []
            sections = []

            # Extract paragraphs
            for para_num, para in enumerate(doc.paragraphs, 1):
                if para.text.strip():
                    text_parts.append(para.text)

            # Extract tables
            for table_num, table in enumerate(doc.tables, 1):
                table_text_parts = []
                for row in table.rows:
                    row_text = '\t'.join(cell.text for cell in row.cells)
                    table_text_parts.append(row_text)

                table_text = '\n'.join(table_text_parts)
                text_parts.append(f"\n[Table {table_num}]\n{table_text}\n")

            full_text = '\n'.join(text_parts)

            sections.append({
                'section_type': 'document',
                'section_title': 'Full Document',
                'text': full_text,
                'section_order': 1,
                'word_count': len(full_text.split()),
                'char_count': len(full_text)
            })

            return {
                'text': full_text,
                'method': 'python-docx',
                'sections': sections,
                'word_count': len(full_text.split()),
                'char_count': len(full_text)
            }

        except Exception as e:
            logger.error(f"Word extraction failed for {file_path.name}: {e}")
            return {
                'text': '',
                'error': str(e)
            }

    # ========================================================================
    # ZIP Extraction
    # ========================================================================

    def extract_zip_text(self, file_path: Path) -> Dict[str, any]:
        """
        Extract text from ZIP archives by:
        1. Extracting all files to a temp directory
        2. Processing each supported file using existing extractors
        3. Combining all extracted text

        Returns: {
            'text': str,
            'method': str,
            'sections': List[Dict],
            'files_processed': int,
            'files_skipped': int
        }
        """
        import zipfile
        import tempfile
        import shutil

        if not zipfile.is_zipfile(str(file_path)):
            return {
                'text': '',
                'error': 'Not a valid ZIP file'
            }

        temp_dir = None
        try:
            # Create temp directory for extraction
            temp_dir = Path(tempfile.mkdtemp(prefix='zip_extract_'))

            # Extract ZIP contents
            with zipfile.ZipFile(str(file_path), 'r') as zf:
                # Check for password protection
                for info in zf.infolist():
                    if info.flag_bits & 0x1:  # Encrypted flag
                        return {
                            'text': '',
                            'error': 'ZIP file is password-protected',
                            'encrypted': True
                        }

                # Extract all files
                try:
                    zf.extractall(temp_dir)
                except Exception as extract_error:
                    return {
                        'text': '',
                        'error': f'Failed to extract ZIP: {extract_error}'
                    }

            # Supported extensions for processing
            supported_extensions = {'.pdf', '.hwp', '.hwpx', '.xlsx', '.xls', '.xlsm', '.docx'}

            # Find all extracted files
            all_files = []
            for root, dirs, files in os.walk(temp_dir):
                for file in files:
                    file_path_extracted = Path(root) / file
                    if file_path_extracted.suffix.lower() in supported_extensions:
                        all_files.append(file_path_extracted)

            if not all_files:
                return {
                    'text': '',
                    'error': 'ZIP contains no supported document types',
                    'sections': []
                }

            # Process each file
            all_sections = []
            all_text_parts = []
            files_processed = 0
            files_skipped = 0
            section_order = 1

            for extracted_file in all_files:
                try:
                    # Use existing extractors based on file type
                    result = self._extract_single_file(extracted_file)

                    if result.get('error'):
                        logger.warning(f"  ⚠️  ZIP inner file {extracted_file.name}: {result['error']}")
                        files_skipped += 1
                        continue

                    if result.get('text'):
                        # Add file header
                        file_header = f"\n{'='*60}\n📄 {extracted_file.name}\n{'='*60}\n"
                        all_text_parts.append(file_header + result['text'])

                        # Add sections with file context
                        for section in result.get('sections', []):
                            section['section_title'] = f"[{extracted_file.name}] {section.get('section_title', '')}"
                            section['section_order'] = section_order
                            section['source_file'] = extracted_file.name
                            all_sections.append(section)
                            section_order += 1

                        files_processed += 1
                        logger.info(f"    ✅ Processed: {extracted_file.name}")

                except Exception as file_error:
                    logger.warning(f"  ⚠️  Error processing {extracted_file.name}: {file_error}")
                    files_skipped += 1

            # Combine all text
            full_text = '\n\n'.join(all_text_parts)

            if not full_text.strip():
                return {
                    'text': '',
                    'error': f'No text extracted from {len(all_files)} files in ZIP',
                    'sections': []
                }

            logger.info(f"  [OK] ZIP: Extracted text from {files_processed} files ({files_skipped} skipped)")

            return {
                'text': full_text,
                'method': 'zip_extraction',
                'sections': all_sections,
                'word_count': len(full_text.split()),
                'char_count': len(full_text),
                'files_processed': files_processed,
                'files_skipped': files_skipped
            }

        except Exception as e:
            logger.error(f"ZIP extraction failed for {file_path.name}: {e}")
            return {
                'text': '',
                'error': str(e)
            }

        finally:
            # Clean up temp directory
            if temp_dir and temp_dir.exists():
                try:
                    shutil.rmtree(temp_dir)
                except Exception as cleanup_error:
                    logger.warning(f"Failed to cleanup temp dir: {cleanup_error}")

    def _extract_single_file(self, file_path: Path) -> Dict[str, any]:
        """
        Extract text from a single file using the appropriate extractor.
        Used by ZIP extraction to process inner files.
        """
        file_ext = file_path.suffix.lower()

        if file_ext == '.pdf':
            return self.extract_pdf_text(file_path)
        elif file_ext == '.hwp':
            return self.extract_hwp_text(file_path)
        elif file_ext == '.hwpx':
            return self.extract_hwpx_text(file_path)
        elif file_ext in ['.xlsx', '.xls', '.xlsm']:
            return self.extract_excel_text(file_path)
        elif file_ext == '.docx':
            return self.extract_word_text(file_path)
        else:
            return {
                'text': '',
                'error': f'Unsupported file type: {file_ext}'
            }

    # ========================================================================
    # Main Extraction Method
    # ========================================================================

    def extract_text(self, file_path: Path) -> Dict[str, any]:
        """
        Extract text from any supported file format
        Auto-detects format based on file extension
        """
        if not file_path.exists():
            return {
                'text': '',
                'error': 'File not found'
            }

        file_ext = file_path.suffix.lower()

        # Route to appropriate extractor
        if file_ext == '.pdf':
            return self.extract_pdf_text(file_path)

        elif file_ext == '.hwp':
            return self.extract_hwp_text(file_path)
        
        elif file_ext == '.hwpx':
            return self.extract_hwpx_text(file_path)

        elif file_ext in ['.xlsx', '.xls', '.xlsm']:
            return self.extract_excel_text(file_path)

        elif file_ext in ['.docx', '.doc']:
            if file_ext == '.docx':
                return self.extract_word_text(file_path)
            else:
                return {
                    'text': '',
                    'error': '.doc format not supported (only .docx)'
                }

        elif file_ext == '.zip':
            return self.extract_zip_text(file_path)

        else:
            return {
                'text': '',
                'error': f'Unsupported file type: {file_ext}'
            }


# ============================================================================
# Stage 3 Integration for ETL Pipeline
# ============================================================================

async def extract_and_save_text(
    db_pool: asyncpg.Pool,
    data_lake_root: Path,
    limit: int = None
):
    """
    Extract text from all unprocessed documents in database
    Integrates with ETL pipeline
    """
    logger.info("=" * 80)
    logger.info("STAGE 3: Text Extraction from Documents")
    logger.info("=" * 80)

    extractor = TextExtractor(data_lake_root)

    # Get unprocessed documents from database
    async with db_pool.acquire() as conn:
        query = """
            SELECT document_id, bid_id, file_path, file_type, document_name
            FROM bid_documents
            WHERE text_extracted = FALSE
            ORDER BY created_at
        """

        if limit:
            query += f" LIMIT {limit}"

        docs = await conn.fetch(query)

    total_docs = len(docs)
    logger.info(f"📄 Found {total_docs} documents to process")

    extracted_count = 0
    error_count = 0

    for idx, doc in enumerate(docs, 1):
        document_id = doc['document_id']
        # Normalize path separators (handle Windows backslashes stored in DB)
        stored_path = doc['file_path'].replace('\\', '/')
        file_path = data_lake_root / stored_path
        file_type = doc['file_type']

        logger.info(f"[{idx}/{total_docs}] Extracting: {doc['document_name']} ({file_type})")

        # Extract text
        result = extractor.extract_text(file_path)

        if result.get('error'):
            logger.warning(f"  ⚠️  {result['error']}")
            error_count += 1
            continue

        # Save text sections to database
        try:
            async with db_pool.acquire() as conn:
                async with conn.transaction():
                    # Save each section
                    for section in result.get('sections', []):
                        # section_id must fit VARCHAR(100): doc_id (max 80) + _sec_ (5) + order (up to 10) = ~95
                        doc_id_truncated = document_id[:80] if len(document_id) > 80 else document_id
                        section_id = f"{doc_id_truncated}_sec_{section['section_order']}"

                        # Sanitize text: remove null bytes (0x00) that break PostgreSQL UTF8
                        section_text = section['text'].replace('\x00', '') if section['text'] else ''
                        section_title = section.get('section_title', '').replace('\x00', '') if section.get('section_title') else ''

                        await conn.execute("""
                            INSERT INTO bid_text_sections (
                                section_id, bid_id, document_id,
                                section_type, section_title,
                                text, cleaned_text, language,
                                page_number, section_order,
                                word_count, char_count
                            ) VALUES ($1, $2, $3, $4, $5, $6, $7, $8, $9, $10, $11, $12)
                            ON CONFLICT (section_id) DO UPDATE SET
                                text = EXCLUDED.text,
                                word_count = EXCLUDED.word_count,
                                char_count = EXCLUDED.char_count
                        """,
                            section_id, doc['bid_id'], document_id,
                            section.get('section_type', 'unknown'),
                            section_title,
                            section_text,
                            section_text,  # TODO: Add cleaning logic
                            'ko',
                            section.get('page_number'),
                            section.get('section_order', 1),
                            section.get('word_count', 0),
                            section.get('char_count', 0)
                        )

                    # Mark document as processed
                    await conn.execute("""
                        UPDATE bid_documents
                        SET text_extracted = TRUE, is_processed = TRUE
                        WHERE document_id = $1
                    """, document_id)

            logger.info(f"  ✅ Extracted {len(result['sections'])} sections, {result.get('word_count', 0)} words")
            extracted_count += 1

        except Exception as e:
            logger.error(f"  ❌ Database error: {e}")
            error_count += 1

    logger.info("=" * 80)
    logger.info(f"✅ Stage 3 Complete: {extracted_count} extracted, {error_count} errors")
    logger.info("=" * 80)

    return {
        "extracted": extracted_count,
        "errors": error_count,
        "total": total_docs
    }
